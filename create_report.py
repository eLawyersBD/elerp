# -*- coding: utf-8 -*-
import docx
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import OxmlElement, parse_xml
from docx.oxml.ns import nsdecls, qn

def create_element(name):
    return OxmlElement(name)

def set_cell_background(cell, fill_hex):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{fill_hex}"/>')
    tcPr.append(shd)

def set_cell_margins(cell, top=100, bottom=100, left=150, right=150):
    tcPr = cell._tc.get_or_add_tcPr()
    tcMar = OxmlElement('w:tcMar')
    for m, val in [('w:top', top), ('w:bottom', bottom), ('w:left', left), ('w:right', right)]:
        node = OxmlElement(m)
        node.set(qn('w:w'), str(val))
        node.set(qn('w:type'), 'dxa')
        tcMar.append(node)
    tcPr.append(tcMar)

def build_document():
    doc = docx.Document()
    
    # Page setup
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11.0)
    section.top_margin = Inches(1.0)
    section.bottom_margin = Inches(1.0)
    section.left_margin = Inches(1.0)
    section.right_margin = Inches(1.0)

    # Color Palette - Premium Blue Brand Theme
    PRIMARY_COLOR = RGBColor(30, 58, 138)    # Deep Navy Blue (#1E3A8A)
    SECONDARY_COLOR = RGBColor(79, 70, 229)  # Indigo (#4F46E5)
    TEXT_COLOR = RGBColor(51, 65, 85)        # Slate Dark (#334155)
    MUTED_COLOR = RGBColor(100, 116, 139)    # Slate Light (#64748B)

    # Styles Setup
    style_normal = doc.styles['Normal']
    font_normal = style_normal.font
    font_normal.name = 'Arial'
    font_normal.size = Pt(11)
    font_normal.color.rgb = TEXT_COLOR

    # Title Style
    title_p = doc.add_paragraph()
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title_p.add_run("ACCOUNTICA Cloud ERP")
    title_run.font.name = 'Arial'
    title_run.font.size = Pt(28)
    title_run.font.bold = True
    title_run.font.color.rgb = PRIMARY_COLOR

    subtitle_p = doc.add_paragraph()
    subtitle_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub_run = subtitle_p.add_run("Comprehensive System Review, Architecture & Marketing Report\n")
    sub_run.font.name = 'Arial'
    sub_run.font.size = Pt(13)
    sub_run.font.italic = True
    sub_run.font.color.rgb = MUTED_COLOR
    
    # Thin divider line
    p_div = doc.add_paragraph()
    p_div.alignment = WD_ALIGN_PARAGRAPH.CENTER
    div_run = p_div.add_run("━" * 50)
    div_run.font.color.rgb = SECONDARY_COLOR
    
    # ── SECTION 1: SYSTEM QUESTIONS & ANSWERS (BENGALI) ──
    h1 = doc.add_paragraph()
    h1_run = h1.add_run("১. সিস্টেম পর্যালোচনা ও মূল প্রশ্নোত্তর (System Q&A)")
    h1_run.font.name = 'Arial'
    h1_run.font.size = Pt(18)
    h1_run.font.bold = True
    h1_run.font.color.rgb = PRIMARY_COLOR
    h1.paragraph_format.space_before = Pt(20)
    h1.paragraph_format.space_after = Pt(10)

    questions = [
        (
            "প্রশ্ন ১: এই ওয়েবসাইট বা ওয়েব অ্যাপটির নাম কি?",
            "উত্তর: এই ওয়েব অ্যাপ্লিকেশনটির নাম হচ্ছে \"ACCOUNTICA Cloud ERP\" (বা ACCOUNTICA Cloud ERP Platform)। এটি একটি অত্যন্ত আধুনিক, সমন্বিত ও ক্লাউড-ভিত্তিক এন্টারপ্রাইজ রিসোর্স প্ল্যানিং (ERP) সমাধান।"
        ),
        (
            "প্রশ্ন ২: এই ওয়েবসাইটটির মূল কাজ কি?",
            "উত্তর: ACCOUNTICA Cloud ERP-এর মূল কাজ হলো একটি প্রতিষ্ঠানের সামগ্রিক ব্যবসায়িক প্রক্রিয়াকে স্বয়ংক্রিয় (Automate) এবং সহজতর করা। এটি ইনভেন্টরি, পারচেজ, সেলস, অ্যাকাউন্টিং (Double-Entry Ledger), এইচআরএমএস (HRMS Attendance/Payroll) এবং প্রাতিষ্ঠানিক যোগাযোগ (Support Desk/Mail Client) একই ড্যাশবোর্ড থেকে রিয়েল-টাইমে পরিচালনা করে।"
        ),
        (
            "প্রশ্ন ৩: এই ওয়েবসাইটটি আপনি কেন ব্যবহার করবেন?",
            "উত্তর: ব্যবসায়িক সমন্বয় ও স্বচ্ছতা আনতে এটি অপরিহার্য। এই প্ল্যাটফর্মটি ব্যবহার করলে আপনার বিভিন্ন বিভাগের কাজ এক জায়গায় চলে আসবে, আর্থিক অনিয়ম দূর হবে, রিয়েল-টাইম প্রফিট/লস ও ব্যালেন্স শীট দেখা যাবে, এবং কর্মচারীদের বেতন-ভাতা-উপস্থিতি স্বয়ংক্রিয়ভাবে শ্রম আইন মেনে হিসাব করা যাবে। ফলে কোনো আলাদা সফটওয়্যারের পেছনে বারবার খরচ করতে হবে না।"
        ),
        (
            "প্রশ্ন ৪: এই ওয়েবসাইট থেকে আপনি কি কি সুবিধা পাবেন?",
            "উত্তর: \n• আর্থিক সঠিকতা: আন্তর্জাতিক মানদণ্ড অনুযায়ী ডাবল-এন্ট্রি অ্যাকাউন্টিং ও ডেবিট/ক্রেডিট ভাউচার ট্র্যাকিং।\n• কাজের সময় ও খরচ হ্রাস: ম্যানুয়াল কাজের ঝামেলা এড়িয়ে এক ক্লিকে ইনভয়েস, পেমেন্ট, ও বেতন হিসাব করার সুবিধা।\n• গতিশীল সিদ্ধান্ত গ্রহণ: ব্যবসায়ের লাভ-ক্ষতি, স্টক ডেটা ও সাপ্লাই চেইন অ্যাক্টিভিটি লাইভ গ্রাফিক্যাল ড্যাশবোর্ডে দেখার সুবিধা।\n• উন্নত গ্রাহক সেবা: বিল্ট-ইন সাপোর্ট টিকেট সিস্টেম ও সিকিউর প্রাতিষ্ঠানিক ইমেইল ম্যানেজমেন্ট।\n• উৎপাদনশীলতা বৃদ্ধি: এআই অ্যাসিস্ট্যান্ট (AI Copilot) এর মাধ্যমে দ্রুত ইমেইল খসড়া ও ক্লায়েন্ট সেন্টিমেন্ট এনালাইসিস।"
        ),
        (
            "প্রশ্ন ৫: এই ওয়েবসাইটে কি কি ফিচার আছে?",
            "উত্তর: ACCOUNTICA Cloud ERP-তে মোট ৩টি মূল সেকশন ও ১০টির বেশি মডিউল রয়েছে:\n১. ইনভেন্টরি ও সাপ্লাই চেইন: প্রোডাক্ট ডেটাবেজ, স্টক মুভমেন্ট হিস্ট্রি, সাপ্লায়ার ডেটাবেজ।\n২. প্রকিউরমেন্ট ও পারচেজ: প্রকিউরমেন্ট রিকুইজিশন, অনুমোদন কাজের ধারা, পারচেজ অর্ডার, সাপ্লায়ার বিলিং।\n৩. সেলস ও কাস্টমার সার্ভিস: কাস্টমার ক্যাটালগ, ড্রাফট ও কনফার্মড অর্ডার, সেলস ইনভয়েস, কাস্টমার লেজার।\n৪. অ্যাকাউন্টিং ও ফাইন্যান্স: জার্নাল/ক্যাশ/ব্যাংক ভাউচার এন্ট্রি, ট্রায়াল ব্যালেন্স, প্রফিট অ্যান্ড লস, ব্যালেন্স শীট।\n৫. এইচআরএমএস ও পে-রোল: জিপিএস জিওফেন্সড অ্যাটেনডেন্স ট্র্যাকিং, লিভ রিকোয়েস্ট ও ব্যালেন্স ট্র্যাকিং, পে-রোল প্রসেসিং, বাংলাদেশ শ্রম আইন অনুযায়ী ট্যাক্স (TDS) ক্যালকুলেটর, এমপ্লয়ী সেলফ সার্ভিস (ESS) পোর্টাল।\n৬. অফিস মেইল ও সাপোর্ট পোর্টাল: IMAP/SMTP ইমেইল ক্লায়েন্ট, কাস্টম ফোল্ডার, রুলস ভিত্তিক ইমেইল রাউটিং, সাপোর্ট টিকিট ম্যানেজমেন্ট, অডিট লগ, একটিভ সেশন ট্র্যাকার, এআই রাইটার।"
        ),
        (
            "প্রশ্ন ৬: ভবিষ্যতে এই ওয়েবসাইটে কোনো ফিচার কাস্টমাইজ করে নেওয়া যাবে কিনা?",
            "উত্তর: হ্যাঁ, অবশ্যই নেওয়া যাবে। এটি একটি অত্যন্ত ফ্লেক্সিবল ও মডুলার আর্কিটেকচার (React, Tailwind CSS, Firebase ও REST APIs) দ্বারা নির্মিত। আপনার প্রতিষ্ঠানের ব্যবসার প্রকৃতি অনুযায়ী যেকোনো ভাউচার ফরম্যাট, ইনভয়েস ডিজাইন, নতুন অনুমোদনের লেভেল অথবা থার্ড-পার্টি ব্যাংক API বা পেমেন্ট গেটওয়ে ইন্টিগ্রেশন করার কাস্টমাইজেশন সুবিধা রয়েছে।"
        ),
        (
            "প্রশ্ন ৭: এই সিস্টেমটি টেস্ট করার কোনো সুযোগ আছে কি?",
            "উত্তর: হ্যাঁ, সম্পূর্ণ ফ্রি টেস্ট করার চমৎকার সুযোগ রয়েছে। টেস্টিং পারপাসে যে কোনো ইউজার কাস্টম ডোমেইনে (erp.elawyersbd.com) গিয়ে গুগল (Google) দিয়ে লগইন করে ৩ ঘণ্টার জন্য গেস্ট হিসেবে সম্পূর্ণ সিস্টেমটি ভিজিট করে দেখতে পারেন এবং এর প্রতিটি ফিচারের কার্যকারিতা নিজে যাচাই করতে পারেন।"
        )
    ]

    for q, a in questions:
        p_q = doc.add_paragraph()
        run_q = p_q.add_run(q)
        run_q.font.bold = True
        run_q.font.color.rgb = SECONDARY_COLOR
        p_q.paragraph_format.space_before = Pt(8)
        p_q.paragraph_format.space_after = Pt(4)

        p_a = doc.add_paragraph()
        p_a.add_run(a)
        p_a.paragraph_format.space_after = Pt(12)

    # ── SECTION 2: MODULE ARCHITECTURE & DETAILED FEATURES ──
    h2 = doc.add_paragraph()
    h2_run = h2.add_run("২. মডিউল-ভিত্তিক বিস্তারিত ফিচার বিবরণী (Detailed Features)")
    h2_run.font.name = 'Arial'
    h2_run.font.size = Pt(18)
    h2_run.font.bold = True
    h2_run.font.color.rgb = PRIMARY_COLOR
    h2.paragraph_format.space_before = Pt(24)
    h2.paragraph_format.space_after = Pt(10)

    modules = [
        ("১. ডাবল-এন্ট্রি অ্যাকাউন্টিং ও ফাইন্যান্স", 
         "আন্তর্জাতিক অ্যাকাউন্টিং স্ট্যান্ডার্ড অনুযায়ী ডাবল-এন্ট্রি সিস্টেম। এখানে ডেবিট/ক্রেডিট ভাউচার এন্ট্রি (Cash Receipt, Cash Payment, Bank, Journal), চার্ট অফ অ্যাকাউন্টস, কাস্টম লেজার বুক, ট্রায়াল ব্যালেন্স, স্বয়ংক্রিয় লাভ-ক্ষতি হিসাবপত্র (Profit & Loss Statement) এবং ব্যালেন্স শীট (Balance Sheet) তৈরির স্বয়ংক্রিয় ব্যবস্থা রয়েছে। রিয়েল-টাইমে ব্যালেন্স শীটের উভয় পাশ স্বয়ংক্রিয়ভাবে মেলানো হয়।"),
        ("২. স্মার্ট ইনভেন্টরি ও সাপ্লাই চেইন", 
         "পণ্য মজুদ রাখার জন্য এটি একটি অনন্য মডিউল। প্রতিটি প্রোডাক্টের বিবরণ, সাপ্লায়ার ডাটাবেজ, স্টক বৃদ্ধি বা হ্রাসের হিসাব (Stock Movement Logs), এবং ন্যূনতম স্টক স্তরের এলার্ট এর অন্তর্ভুক্ত। এটি সাপ্লাই চেইনের অপচয় প্রতিরোধ করে সঠিক স্টক কাউন্ট ধরে রাখে।"),
        ("৩. ক্রয় ও প্রকিউরমেন্ট প্ল্যানিং", 
         "প্রতিষ্ঠান বা ফ্যাক্টরির জন্য রিকুইজিশন তৈরি থেকে শুরু করে সরবরাহকারী নির্বাচন, বাজেট চেক, অনুমোদন প্রবাহ (Approval Workflow), পারচেজ অর্ডার (Purchase Order) জেনারেশন এবং বিল ভাউচার এন্ট্রি এর মাধ্যমে সম্পূর্ণ ক্রয়ের প্রক্রিয়াকে ট্র্যাকিং ও নিয়ন্ত্রণ করা সম্ভব।"),
        ("৪. বিক্রয় ও রেভিনিউ ট্র্যাকিং", 
         "কাস্টমার প্রোফাইল তৈরি, পণ্যের প্রাইসিং ক্যাটালগ, ড্রাফট ইনভয়েস তৈরি, ডিসকাউন্ট ও ভ্যাট হিসাব করে ফাইনাল সেলস ইনভয়েস তৈরি করা এবং পেমেন্ট রিসিভ ও বাকি টাকা ট্র্যাকিংয়ের সমন্বিত মডিউল। কাস্টমার লেজার স্বয়ংক্রিয়ভাবে সেলস ডেটার সাথে যুক্ত।"),
        ("৫. এইচআরএমএস, উপস্থিতি ও পে-রোল (ESS)", 
         "এটি একটি অত্যন্ত সমৃদ্ধ মডিউল। কর্মচারীদের তথ্য সংরক্ষণের ডিরেক্টরি, জিপিএস ও জিওফেন্সিং (Geofencing) প্রযুক্তির মাধ্যমে হেড অফিস বা যেকোনো ব্রাঞ্চ থেকে উপস্থিতি রেকর্ড করার সিস্টেম (Attendance Logs), ছুটির আবেদন প্রসেস করা, এক ক্লিকে বেতন প্রস্তুত করা (Payroll processing), এবং বাংলাদেশ ট্যাক্স রুলস (Basic, House rent, Conveyance, Sanchaypatra Investment) অনুযায়ী ট্যাক্স ও TDS হিসাব করার নির্ভরযোগ্য ক্যালকুলেটর। কর্মচারীদের জন্য রয়েছে নিজস্ব সেলফ সার্ভিস (ESS) পোর্টাল।"),
        ("৬. অফিস ইমেইল ও সাপোর্ট পোর্টাল", 
         "এটি জিমেইল, আউটলুক বা যেকোনো নিজস্ব ডোমেইনের মেলবক্সের সাথে IMAP/SMTP প্রোটোকলে যুক্ত হয়। প্রাতিষ্ঠানিক যোগাযোগ সহজ করার জন্য ইনবক্স, সেন্ট, ড্রাফট, ট্র্যাশ ফোল্ডারের পাশাপাশি নতুন ফোল্ডার তৈরি, অটোমেটিক ফিল্টার রুলস সেটআপ করা এবং প্রতিটি ইনকামিং ইমেইলকে কাস্টমার সাপোর্ট টিকেট (Support Ticket System) হিসেবে রূপান্তর করে নির্দিষ্ট এজেন্টের কাছে অ্যাসাইন করার জন্য টিকেট ম্যানেজমেন্ট বোর্ড রয়েছে। ইমেইল খসড়া লিখতে এবং সেন্টিমেন্ট বুঝতে সহায়তা করে বিল্ট-ইন জেনারেটিভ AI Assistant।")
    ]

    for title, desc in modules:
        p_mt = doc.add_paragraph()
        run_mt = p_mt.add_run(title)
        run_mt.font.bold = True
        run_mt.font.size = Pt(13)
        run_mt.font.color.rgb = SECONDARY_COLOR
        p_mt.paragraph_format.space_before = Pt(8)
        p_mt.paragraph_format.space_after = Pt(3)

        p_md = doc.add_paragraph()
        p_md.add_run(desc)
        p_md.paragraph_format.space_after = Pt(10)

    # ── SECTION 3: MARKETING MATERIALS (BENGALI & ENGLISH) ──
    h3 = doc.add_paragraph()
    h3_run = h3.add_run("৩. মার্কেটিং ও প্রমোশনাল ম্যাটেরিয়ালস (Marketing & Promo Content)")
    h3_run.font.name = 'Arial'
    h3_run.font.size = Pt(18)
    h3_run.font.bold = True
    h3_run.font.color.rgb = PRIMARY_COLOR
    h3.paragraph_format.space_before = Pt(24)
    h3.paragraph_format.space_after = Pt(10)

    # Sub-heading 1
    sh1 = doc.add_paragraph()
    sh1_run = sh1.add_run("১. ফেসবুক ও সোশ্যাল মিডিয়া পোস্ট (বাংলা)")
    sh1_run.font.bold = True
    sh1_run.font.size = Pt(14)
    sh1_run.font.color.rgb = SECONDARY_COLOR
    sh1.paragraph_format.space_after = Pt(6)

    fb_bangla_text = (
        "🚀 আপনার ব্যবসায়ের সকল কার্যক্রম কি এখনো এলোমেলো? আলাদা সফটওয়্যারের পেছনে কি প্রতি মাসে হাজার হাজার টাকা নষ্ট হচ্ছে?\n\n"
        "সব ঝামেলার অবসান ঘটাতে চলে এলো ACCOUNTICA Cloud ERP! 💻✨\n\n"
        "ইনভেন্টরি ট্র্যাকিং থেকে শুরু করে সঠিক ফাইন্যান্সিয়াল অডিট, ডাবল-এন্ট্রি অ্যাকাউন্টিং, জিপিএস অ্যাটেনডেন্স ও পে-রোল সমৃদ্ধ HRMS এবং প্রাতিষ্ঠানিক যোগাযোগে AI Assistant সম্বলিত অফিস ইমেইল পোর্টাল—সব মিলবে একটি প্ল্যাটফর্মেই।\n\n"
        "💼 ACCOUNTICA Cloud ERP এর বিশেষত্বসমূহ:\n"
        "✅ রিয়েল-টাইম ডাবল-এন্ট্রি অ্যাকাউন্টিং (ভাউচার, লেজার, ট্রায়াল ব্যালেন্স, প্রফিট/লস ও ব্যালেন্স শীট)\n"
        "✅ জিপিএস জিওফেন্সড অ্যাটেনডেন্স এবং স্বয়ংক্রিয় পে-রোল (ট্যাক্স হিসাবসহ)\n"
        "✅ স্মার্ট প্রকিউরমেন্ট এবং ইনভেন্টরি ট্র্যাকিং (স্টক মুভমেন্ট এলার্ট)\n"
        "✅ সাপোর্ট ডেসক এবং IMAP/SMTP অফিস ইমেইল উইন্ডো উইথ AI Copilot\n\n"
        "আপনার ব্যবসাকে শতভাগ স্মার্ট ও ডিজিটাল করতে আজই ভিজিট করুন: erp.elawyersbd.com 🌐\n\n"
        "#BusinessAutomation #CloudERP #AccounticaERP #AccountingSoftware #HRMS #SmartBusiness"
    )
    p_fb = doc.add_paragraph()
    p_fb.add_run(fb_bangla_text)
    p_fb.paragraph_format.space_after = Pt(14)

    # Sub-heading 2
    sh2 = doc.add_paragraph()
    sh2_run = sh2.add_run("2. Social Media & Professional Posting (English)")
    sh2_run.font.bold = True
    sh2_run.font.size = Pt(14)
    sh2_run.font.color.rgb = SECONDARY_COLOR
    sh2.paragraph_format.space_after = Pt(6)

    fb_english_text = (
        "🚀 Streamline your entire enterprise with ACCOUNTICA Cloud ERP! Say goodbye to scattered files and manual errors.\n\n"
        "Manage Inventory, Procurement, Sales, Double-Entry Accounting, Geofenced HRMS & Payroll, and secure Organizational Mail client with integrated support ticketing system—all working cohesively in a single cloud dashboard. 💻✨\n\n"
        "💼 Why Choose ACCOUNTICA ERP?\n"
        "✅ Comprehensive Financial Suite: Auto-generate Trial Balance, Profit & Loss sheets, and Balance Sheet.\n"
        "✅ Smart HRMS: Geofenced attendance tracking (HQ & branches), leaves, and automated TDS tax calculator.\n"
        "✅ Procurement & Sales Workflow: Maintain perfect catalogs, purchase approvals, invoicing, and customer ledgers.\n"
        "✅ Office Email & Customer Tickets: Access integrated IMAP/SMTP client with AI Copilot for draft writing & security audits.\n\n"
        "Take the smart step towards automation today. Explore live at: erp.elawyersbd.com 🌐\n\n"
        "#EnterpriseAutomation #CloudERP #Fintech #HRMS #InventoryManagement #AIIntegration"
    )
    p_fbe = doc.add_paragraph()
    p_fbe.add_run(fb_english_text)
    p_fbe.paragraph_format.space_after = Pt(14)

    # Sub-heading 3: Short SMS Campaign
    sh3 = doc.add_paragraph()
    sh3_run = sh3.add_run("৩. ছোট এসএমএস ক্যাম্পেইন (SMS & WhatsApp Campaign)")
    sh3_run.font.bold = True
    sh3_run.font.size = Pt(14)
    sh3_run.font.color.rgb = SECONDARY_COLOR
    sh3.paragraph_format.space_after = Pt(6)

    sms_text = (
        "💬 SMS Template:\n"
        "\"ব্যবসায়ের ইনভেন্টরি, ডাবল-এন্ট্রি অ্যাকাউন্টিং, জিপিএস অ্যাটেনডেন্স এবং অফিস ইমেইল এক প্ল্যাটফর্মে! আজই ব্যবহার করুন ACCOUNTICA Cloud ERP। ভিজিট: erp.elawyersbd.com\"\n\n"
        "💬 WhatsApp Broadcast:\n"
        "\"সম্মানিত ব্যবসায়ী ভাই ও বোনেরা, আপনার ব্যবসায়ের হিসাব-নিকাশ, কর্মচারী ব্যবস্থাপনা এবং ইমেইল সিস্টেমকে স্বয়ংক্রিয় করতে ACCOUNTICA Cloud ERP নিয়ে এসেছে সেরা সমাধান। কোনো আলাদা সার্ভার খরচ ছাড়াই রিয়েল-টাইম ক্লাউড ডেটা অ্যাক্সেস পান। ফ্রী ডেমো দেখতে ভিজিট করুন: erp.elawyersbd.com\""
    )
    p_sms = doc.add_paragraph()
    p_sms.add_run(sms_text)
    p_sms.paragraph_format.space_after = Pt(14)

    # ── SECTION 4: ARCHITECTURE MATRIX ──
    h4 = doc.add_paragraph()
    h4_run = h4.add_run("৪. সফটওয়্যার আর্কিটেকচার ম্যাট্রিক্স (Architecture Matrix)")
    h4_run.font.name = 'Arial'
    h4_run.font.size = Pt(18)
    h4_run.font.bold = True
    h4_run.font.color.rgb = PRIMARY_COLOR
    h4.paragraph_format.space_before = Pt(24)
    h4.paragraph_format.space_after = Pt(10)

    # Add Table
    table = doc.add_table(rows=1, cols=3)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr_cells = table.rows[0].cells
    headers = ["মডিউলের নাম (Module)", "ব্যবহৃত প্রযুক্তি (Stack)", "প্রধান কাজ (Core Purpose)"]
    
    for i, h in enumerate(headers):
        hdr_cells[i].text = h
        set_cell_background(hdr_cells[i], "1E3A8B")
        set_cell_margins(hdr_cells[i], top=120, bottom=120, left=150, right=150)
        p = hdr_cells[i].paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.runs[0]
        run.font.bold = True
        run.font.size = Pt(10)
        run.font.color.rgb = RGBColor(255, 255, 255)

    data = [
        ("Inventory & Procurement", "React, LocalStorage / API", "পণ্য মজুদ, স্টক লেনদেন ইতিহাস, রিকুইজিশন ও সরবরাহকারী বিলিং"),
        ("Accounting & Vouchers", "React Context, LocalStorage / API", "জার্নাল, ব্যাংক/ক্যাশ ভাউচার এন্ট্রি, লেজার বুক, ট্রায়াল ব্যালেন্স, ব্যালেন্স শীট"),
        ("HRMS & Payroll Portal", "GPS Geolocation API, LocalStorage", "উপস্থিতি ট্র্যাকিং, ছুটির আবেদন, শ্রম আইন মেনে বেতন ও ট্যাক্স ক্যালকুলেটর"),
        ("Email Client & Support", "IMAP/SMTP Server Protocols, AI API", "প্রাতিষ্ঠানিক ইমেইল ক্লায়েন্ট, কাস্টমার সাপোর্ট টিকিট, এআই ড্রাফট রাইটার"),
        ("Core Backend Integration", "Firebase Firestore & Firebase Auth", "রিয়েল-টাইম ক্লাউড ডেটা সিঙ্ক ও বিশ্বস্ত ইউজার অথেনটিকেশন ও পারমিশন সিস্টেম")
    ]

    for row_idx, (mod, tech, purp) in enumerate(data):
        row_cells = table.add_row().cells
        row_cells[0].text = mod
        row_cells[1].text = tech
        row_cells[2].text = purp
        
        # Color alternating rows
        bg_color = "F1F5F9" if row_idx % 2 == 0 else "FFFFFF"
        
        for cell in row_cells:
            set_cell_background(cell, bg_color)
            set_cell_margins(cell, top=100, bottom=100, left=150, right=150)
            p = cell.paragraphs[0]
            for run in p.runs:
                run.font.size = Pt(9.5)
                run.font.color.rgb = TEXT_COLOR

    # Footer note
    doc.add_paragraph("\n\n")
    p_foot = doc.add_paragraph()
    p_foot.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    foot_run = p_foot.add_run("ACCOUNTICA Cloud ERP Team © 2026")
    foot_run.font.size = Pt(9.5)
    foot_run.font.italic = True
    foot_run.font.color.rgb = MUTED_COLOR

    doc.save("ACCOUNTICA_ERP_Review_Report_v2.docx")
    print("Successfully generated ACCOUNTICA_ERP_Review_Report_v2.docx")

if __name__ == '__main__':
    build_document()
